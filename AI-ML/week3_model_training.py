"""
Week 3 — Model Training & Evaluation
=====================================
AI-Powered Cybersecurity Threat Detection System
AI/ML Team

What this script does:
  1. Loads cleaned training/test data from CLEAN/ folder
  2. Trains a Random Forest Classifier
  3. Makes predictions on test set
  4. Calculates evaluation metrics (Accuracy, Precision, Recall, F1, Confusion Matrix)
  5. Saves the trained model as random_forest_model.pkl
  6. Generates a detailed evaluation report (Word document)

Requirements:
    pip install pandas numpy scikit-learn joblib matplotlib seaborn python-docx

Usage:
    python week3_model_training.py
"""

import os
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
import joblib
from datetime import datetime

# For report generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ─────────────────────────────────────────────────────────────────────
DATA_DIR      = "CLEAN"
MODEL_FILE    = os.path.join(DATA_DIR, "random_forest_model.pkl")
REPORT_FILE   = "Week3_Model_Evaluation_Report.docx"

# ── Hyperparameters ───────────────────────────────────────────────────────────
RF_PARAMS = {
    "n_estimators": 100,
    "max_depth": 20,
    "min_samples_split": 10,
    "min_samples_leaf": 4,
    "random_state": 42,
    "n_jobs": -1,  # use all CPU cores
    "verbose": 1,
}


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Load cleaned data
# ══════════════════════════════════════════════════════════════════════════════
def load_data(data_dir):
    print(f"\n{'='*55}")
    print("STEP 1 — Loading cleaned data from CLEAN/")
    print(f"{'='*55}")

    X_train = pd.read_csv(os.path.join(data_dir, "X_train.csv"))
    y_train = pd.read_csv(os.path.join(data_dir, "y_train.csv")).squeeze()
    X_test = pd.read_csv(os.path.join(data_dir, "X_test.csv"))
    y_test = pd.read_csv(os.path.join(data_dir, "y_test.csv")).squeeze()

    print(f"  X_train: {X_train.shape[0]:,} rows × {X_train.shape[1]} features")
    print(f"  y_train: {y_train.shape[0]:,} labels")
    print(f"  X_test:  {X_test.shape[0]:,} rows × {X_test.shape[1]} features")
    print(f"  y_test:  {y_test.shape[0]:,} labels")
    print(f"  Classes: {sorted(y_test.unique())}")

    return X_train, y_train, X_test, y_test


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Train Random Forest model
# ══════════════════════════════════════════════════════════════════════════════
def train_model(X_train, y_train, params):
    print(f"\n{'='*55}")
    print("STEP 2 — Training Random Forest Classifier")
    print(f"{'='*55}")
    print(f"  Hyperparameters:")
    for key, val in params.items():
        if key != "verbose" and key != "n_jobs":
            print(f"    {key}: {val}")

    model = RandomForestClassifier(**params)
    model.fit(X_train, y_train)

    print(f"\n  ✅ Model training complete!")
    return model


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — Make predictions
# ══════════════════════════════════════════════════════════════════════════════
def predict(model, X_test):
    print(f"\n{'='*55}")
    print("STEP 3 — Making predictions on test set")
    print(f"{'='*55}")

    y_pred = model.predict(X_test)
    print(f"  Predicted {len(y_pred):,} test samples")
    return y_pred


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Evaluate model
# ══════════════════════════════════════════════════════════════════════════════
def evaluate_model(y_test, y_pred, label_encoder):
    print(f"\n{'='*55}")
    print("STEP 4 — Evaluating model performance")
    print(f"{'='*55}")

    # Overall accuracy
    accuracy = accuracy_score(y_test, y_pred)
    print(f"\n  Overall Accuracy: {accuracy:.4f} ({accuracy*100:.2f}%)")

    # Per-class metrics
    print(f"\n  Per-Class Metrics:")
    print(f"  {'Attack Type':<20} {'Precision':<12} {'Recall':<12} {'F1-Score':<12}")
    print(f"  {'-'*56}")

    precision = precision_score(y_test, y_pred, average=None, zero_division=0)
    recall = recall_score(y_test, y_pred, average=None, zero_division=0)
    f1 = f1_score(y_test, y_pred, average=None, zero_division=0)

    class_names = label_encoder.classes_
    for i, class_name in enumerate(class_names):
        print(f"  {class_name:<20} {precision[i]:<12.4f} {recall[i]:<12.4f} {f1[i]:<12.4f}")

    # Confusion matrix
    cm = confusion_matrix(y_test, y_pred)

    return {
        "accuracy": accuracy,
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "confusion_matrix": cm,
        "class_names": class_names,
        "classification_report": classification_report(y_test, y_pred, target_names=class_names, digits=4)
    }


# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — Save model
# ══════════════════════════════════════════════════════════════════════════════
def save_model(model, model_file):
    print(f"\n{'='*55}")
    print("STEP 5 — Saving trained model")
    print(f"{'='*55}")

    joblib.dump(model, model_file)
    print(f"  Model saved → {model_file}")


# ══════════════════════════════════════════════════════════════════════════════
# STEP 6 — Generate evaluation report
# ══════════════════════════════════════════════════════════════════════════════
def generate_report(model, metrics, X_train, y_train, X_test, y_test, report_file):
    print(f"\n{'='*55}")
    print("STEP 6 — Generating evaluation report")
    print(f"{'='*55}")

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI-Powered Cybersecurity Threat Detection System")
    title_run.font.size = Pt(24)
    title_run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Week 3 — Model Training & Evaluation Report")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(46, 116, 181)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_paragraph()  # gap

    # Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "This report documents the training and evaluation of a Random Forest Classifier "
        "trained on the CICIDS2017 dataset to detect cybersecurity threats. The model was trained "
        "on 2,014,924 samples and evaluated on 503,731 held-out test samples."
    )

    # Model Configuration
    doc.add_heading("2. Model Configuration", level=1)
    doc.add_paragraph("Random Forest Classifier with the following hyperparameters:")
    for key, val in RF_PARAMS.items():
        if key not in ["verbose", "n_jobs"]:
            doc.add_paragraph(f"{key}: {val}", style="List Bullet")

    # Training Data
    doc.add_heading("3. Training Data", level=1)
    doc.add_paragraph(f"Training samples: {X_train.shape[0]:,}")
    doc.add_paragraph(f"Test samples: {X_test.shape[0]:,}")
    doc.add_paragraph(f"Features: {X_train.shape[1]}")
    doc.add_paragraph(f"Attack classes: {len(metrics['class_names'])}")

    # Results
    doc.add_heading("4. Evaluation Results", level=1)

    # Overall accuracy
    doc.add_heading("4.1 Overall Accuracy", level=2)
    accuracy_para = doc.add_paragraph()
    accuracy_run = accuracy_para.add_run(f"{metrics['accuracy']:.4f} ({metrics['accuracy']*100:.2f}%)")
    accuracy_run.font.bold = True
    accuracy_run.font.size = Pt(14)

    # Per-class metrics table
    doc.add_heading("4.2 Per-Class Metrics", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Attack Type"
    hdr_cells[1].text = "Precision"
    hdr_cells[2].text = "Recall"
    hdr_cells[3].text = "F1-Score"

    for i, class_name in enumerate(metrics['class_names']):
        row_cells = table.add_row().cells
        row_cells[0].text = class_name
        row_cells[1].text = f"{metrics['precision'][i]:.4f}"
        row_cells[2].text = f"{metrics['recall'][i]:.4f}"
        row_cells[3].text = f"{metrics['f1'][i]:.4f}"

    # Confusion Matrix
    doc.add_heading("4.3 Confusion Matrix", level=2)
    cm = metrics['confusion_matrix']
    cm_table = doc.add_table(rows=len(metrics['class_names'])+1, cols=len(metrics['class_names'])+1)
    cm_table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = cm_table.rows[0].cells
    header_cells[0].text = "Predicted →"
    for j, class_name in enumerate(metrics['class_names']):
        header_cells[j+1].text = class_name[:8]  # abbreviate

    # Data rows
    for i, class_name in enumerate(metrics['class_names']):
        row_cells = cm_table.rows[i+1].cells
        row_cells[0].text = class_name[:8]
        for j in range(len(metrics['class_names'])):
            row_cells[j+1].text = str(cm[i, j])

    # Classification report
    doc.add_heading("4.4 Detailed Classification Report", level=2)
    report_para = doc.add_paragraph(metrics['classification_report'], style='List')
    report_para.runs[0].font.name = 'Courier New'
    report_para.runs[0].font.size = Pt(9)

    # Interpretation
    doc.add_heading("5. Interpretation", level=1)
    doc.add_paragraph(
        "The confusion matrix shows predicted vs actual attack types. Diagonal elements represent "
        "correct predictions. Off-diagonal elements show misclassifications. Review high off-diagonal "
        "values to identify which attack types are commonly confused."
    )

    # Conclusion
    doc.add_heading("6. Conclusion", level=1)
    accuracy_pct = metrics['accuracy'] * 100
    if accuracy_pct >= 95:
        assessment = "Excellent"
    elif accuracy_pct >= 90:
        assessment = "Very Good"
    elif accuracy_pct >= 85:
        assessment = "Good"
    else:
        assessment = "Acceptable (may need tuning)"

    doc.add_paragraph(
        f"The Random Forest model achieves {accuracy_pct:.2f}% accuracy on the test set, indicating "
        f"{assessment} performance in classifying cybersecurity threats. The model is ready for "
        "integration with the FastAPI backend in Week 5."
    )

    # Save
    doc.save(report_file)
    print(f"  Report saved → {report_file}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n" + "="*55)
    print("  Week 3 — Model Training & Evaluation")
    print("  CICIDS2017 | AI/ML Cybersecurity Project")
    print("="*55)

    # Load data
    X_train, y_train, X_test, y_test = load_data(DATA_DIR)

    # Load label encoder (created in Week 2)
    label_encoder = joblib.load(os.path.join(DATA_DIR, "label_encoder.pkl"))

    # Train model
    model = train_model(X_train, y_train, RF_PARAMS)

    # Predict
    y_pred = predict(model, X_test)

    # Evaluate
    metrics = evaluate_model(y_test, y_pred, label_encoder)

    # Save model
    save_model(model, MODEL_FILE)

    # Generate report
    generate_report(model, metrics, X_train, y_train, X_test, y_test, REPORT_FILE)

    print(f"\n{'='*55}")
    print("  ✅ Week 3 Complete!")
    print(f"  Model:  {MODEL_FILE}")
    print(f"  Report: {REPORT_FILE}")
    print(f"{'='*55}\n")
